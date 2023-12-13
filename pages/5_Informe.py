import docx
import streamlit as st
import pandas as pd
import numpy as np
from math import pi

st.set_page_config(
    page_title = 'Calculadora de gálibos',
    page_icon = '🚊',
    layout = 'centered'
)

st.markdown ("# Informe")

if st.session_state.ancho_via == 1.668:
    ancho_via = "Ibérico (1.668 mm)"
elif st.session_state.ancho_via == 1.435:
    ancho_via = "Europeo (1.435 mm)"
elif st.session_state.ancho_via == 1.000:
    ancho_via = "Métrico (1.000 mm)"

st.data_editor(st.session_state.input, column_config = {'Peralte máximo (mm)': st.column_config.Column(width = 'small', required = True),
'Insuficiencia de peralte máximo (mm)': st.column_config.Column(width = 'small', required = True),
'Radio mínimo en planta (m)': st.column_config.Column(width = 'small', required = True),
'Radio mínimo acuerdo vertical (m)': st.column_config.Column(width = 'small', required = True),
'Velocidad máxima de circulación (km/h)': st.column_config.Column(width = 'small', required = True)})

if st.button("Variables"):
    st.write("El ancho de vía considerado ha sido **%s**, vía **%s** en **%s**" % (ancho_via, st.session_state.tipo_linea, st.session_state.tipo_via))
    st.write("A continuación, se muestran las condiciones de trazado consideradas:")
    st.dataframe(st.session_state.input, hide_index = True, use_container_width = True)
    st.write("El gálibo considerado para el cálculo ha sido **%s** para las partes altas y **%s** para las partes bajas." %
    (st.session_state.partes_altas, st.session_state.partes_bajas))
    st.table(st.session_state.var)
    st.data_editor(st.session_state.var, column_config = {'name': st.latex(r'''α_{c} (º)''')})

st.latex('''
    \begin{table}
    \begin{tabular}{ |c|c|c| } 
    \hline
    cell1 & cell2 & cell3 \\ 
    cell4 & cell5 & cell6 \\ 
    cell7 & cell8 & cell9 \\ 
    \hline
    \end{tabular}
    \end{table}''')

if st.button("Generar informe"):
    # Create a document
    doc = docx.Document()

    # Add a paragraph to the document
    p = doc.add_paragraph()

    # Add some formatting to the paragraph
    p.paragraph_format.line_spacing = 1
    p.paragraph_format.space_after = 0

    # Add a run to the paragraph
    run = p.add_run("python-docx")

    # Add some formatting to the run
    run.bold = True
    run.italic = True
    run.font.name = 'Arial'
    run.font.size = docx.shared.Pt(16)

    # Add more text to the same paragraph
    run = p.add_run(" Tutorial")

    # Format the run
    run.bold = True
    run.font.name = 'Arial'
    run.font.size = docx.shared.Pt(16)

    # Add another paragraph (left blank for an empty line)
    doc.add_paragraph()

    # Add another paragraph
    p = doc.add_paragraph()

    # Add a run and format it
    run = p.add_run("El valor del gálibo límite es... [bla bla bla]")
    run.font.name = 'Arial'
    run.font.size = docx.shared.Pt(12)

    # Save the document
    doc.save("docx-python Tutorial Demo.docx")